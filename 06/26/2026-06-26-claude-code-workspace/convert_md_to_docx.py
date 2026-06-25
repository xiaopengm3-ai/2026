import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

with open(r'C:\Users\Administrator\Desktop\鹿隐LUYIN-产品介绍与销售方案.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

def add_styled_paragraph(doc, text, size=None, bold=False, color=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if size: run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)
    if alignment is not None: p.alignment = alignment
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_formatted_paragraph(doc, text, size=11):
    """Parse inline markdown: **bold**, *italic*"""
    p = doc.add_paragraph()
    # Split by bold markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(size)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            run = p.add_run(part)
            run.font.size = Pt(size)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    # H1
    if line.startswith('# ') and not line.startswith('## '):
        text = line[2:]
        p = add_styled_paragraph(doc, text, size=22, bold=True, color=(90, 58, 32))
        p.space_before = Pt(24)
        p.space_after = Pt(12)

    # H2
    elif line.startswith('## '):
        text = line[3:]
        p = add_styled_paragraph(doc, text, size=16, bold=True, color=(90, 58, 32))
        p.space_before = Pt(18)
        p.space_after = Pt(8)

    # H3
    elif line.startswith('### '):
        text = line[4:]
        p = add_styled_paragraph(doc, text, size=13, bold=True, color=(74, 106, 58))
        p.space_before = Pt(14)
        p.space_after = Pt(6)

    # H4
    elif line.startswith('#### '):
        text = line[5:]
        p = add_styled_paragraph(doc, text, size=12, bold=True)
        p.space_before = Pt(10)
        p.space_after = Pt(4)

    # Horizontal rule
    elif line.strip() == '---':
        doc.add_paragraph('─' * 60)

    # Table
    elif line.startswith('|') and i + 2 < len(lines) and lines[i+1].startswith('|') and '---' in lines[i+1]:
        # Collect table lines
        table_lines = [line]
        i += 1
        while i < len(lines) and lines[i].startswith('|'):
            table_lines.append(lines[i])
            i += 1
        i -= 1  # will be incremented at end of loop

        # Skip separator line
        header_line = table_lines[0]
        data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|$', l) and l != header_line]
        if not data_lines:
            data_lines = table_lines[2:] if len(table_lines) > 2 else []

        # Parse cells
        all_lines = [header_line] + data_lines
        cells_list = []
        for tl in all_lines:
            cells = [c.strip() for c in tl.split('|')[1:-1]]
            cells_list.append(cells)

        max_cols = max(len(c) for c in cells_list)

        table = doc.add_table(rows=len(cells_list), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, cells in enumerate(cells_list):
            for col_idx, cell_text in enumerate(cells):
                if col_idx < max_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if row_idx == 0:
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if row_idx == 0:
                        # Header row shading
                        shading = cell._element.get_or_add_tcPr()

        doc.add_paragraph()  # space after table

    # Bullet list
    elif line.startswith('- ') or line.startswith('  - '):
        indent_level = 0 if line.startswith('- ') else 1
        text = line.lstrip('- ').strip()
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Ordered list
    elif re.match(r'^\d+\. ', line):
        text = re.sub(r'^\d+\. ', '', line)
        p = doc.add_paragraph(style='List Number')
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Bold text line (like **消费者**：xxx)
    elif line.startswith('**') and '：' in line:
        add_formatted_paragraph(doc, line, size=11)

    # Quote / block
    elif line.startswith('> '):
        text = line[2:]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Empty line
    elif line.strip() == '':
        pass  # skip, docx handles spacing

    # Normal paragraph
    else:
        add_formatted_paragraph(doc, line, size=11)

    i += 1

output_path = r'C:\Users\Administrator\Desktop\鹿隐LUYIN-产品介绍与销售方案.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
