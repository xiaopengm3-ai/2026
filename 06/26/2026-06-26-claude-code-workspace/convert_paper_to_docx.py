import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── Default style ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.25
pf.first_line_indent = Cm(0.74)  # ~2 chars for 10.5pt
pf.space_before = Pt(0)
pf.space_after = Pt(0)

def add_run_to_paragraph(p, text, size, bold, color, font_west, font_east, italic=False):
    """Add a single run with full formatting."""
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = font_west
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_east)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_p(doc, text, size, bold, color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=0, space_after=0, indent=False, font_west='Times New Roman', font_east='宋体'):
    """Add a styled paragraph for headings."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    add_run_to_paragraph(p, text, size, bold, color, font_west, font_east)
    return p

def add_body_p(doc, text, size=10.5, bold=False, color=None, font_west='Times New Roman', font_east='宋体',
               indent=True, italic=False):
    """Add a body paragraph with first-line indent."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Parse inline **bold** markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run_to_paragraph(p, part[2:-2], size, True, color, font_west, font_east, italic)
        else:
            add_run_to_paragraph(p, part, size, bold, color, font_west, font_east, italic)
    return p

def add_ref_p(doc, text, size=12, font_west='Times New Roman', font_east='宋体'):
    """Add a reference entry."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    # Parse [N] label
    m = re.match(r'^(\[\d+\])\s*(.*)', text)
    if m:
        add_run_to_paragraph(p, m.group(1) + ' ', size, False, None, font_west, font_east)
        add_run_to_paragraph(p, m.group(2), size, False, None, font_west, font_east)
    else:
        add_run_to_paragraph(p, text, size, False, None, font_west, font_east)
    return p

# ── Read markdown ──
md_path = r'e:\claude code files\地缘政治压力下中国数字平台企业的国际市场进入模式动态调整研究.md'
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_refs = False
ref_entries = []

while i < len(lines):
    line = lines[i].rstrip()

    # Detect references section
    if line.startswith('## 参考文献'):
        in_refs = True
        # Print ref heading
        add_heading_p(doc, '参考文献', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=18, space_after=12)
        i += 1
        continue

    if in_refs:
        if line.strip() == '':
            i += 1
            continue
        if re.match(r'^\[\d+\]', line):
            add_ref_p(doc, line, size=12)
        i += 1
        continue

    # H1 — paper title (三号 16pt)
    if line.startswith('# ') and not line.startswith('## '):
        text = line[2:]
        add_heading_p(doc, text, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=6, space_after=18)

    # H2 — 一级标题 (小三 15pt)
    elif line.startswith('## '):
        text = line[3:]
        add_heading_p(doc, text, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=16, space_after=8, indent=False)

    # H3 — third-level heading (四号 14pt)
    elif line.startswith('### '):
        text = line[4:]
        add_heading_p(doc, text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=12, space_after=6, indent=False)

    # H4 — fourth-level heading (小四 12pt)
    elif line.startswith('#### '):
        text = line[5:]
        add_heading_p(doc, text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=10, space_after=4, indent=False)

    # Horizontal rule
    elif line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        add_run_to_paragraph(p, '─' * 60, 10.5, False, None, 'Times New Roman', '宋体')

    # Table
    elif line.startswith('|') and i + 1 < len(lines) and lines[i+1].startswith('|') and '---' in lines[i+1]:
        table_lines = [line]
        i += 1
        while i < len(lines) and lines[i].startswith('|'):
            table_lines.append(lines[i])
            i += 1
        i -= 1

        # Parse
        header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
        data_cells_list = []
        for tl in table_lines[2:]:
            cells = [c.strip() for c in tl.split('|')[1:-1]]
            data_cells_list.append(cells)

        max_cols = max(len(header_cells), max((len(d) for d in data_cells_list), default=0))
        from docx.enum.table import WD_TABLE_ALIGNMENT
        table = doc.add_table(rows=1 + len(data_cells_list), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for ci, ct in enumerate(header_cells):
            cell = table.cell(0, ci)
            cell.text = ''
            run = cell.paragraphs[0].add_run(ct)
            run.font.size = Pt(10)
            run.bold = True
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for ri, cells in enumerate(data_cells_list):
            for ci, ct in enumerate(cells):
                if ci < max_cols:
                    cell = table.cell(ri + 1, ci)
                    cell.text = ''
                    run = cell.paragraphs[0].add_run(ct)
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_paragraph()

    # Bold heading-like inline (e.g. "**H1：** ...")
    elif re.match(r'^\*\*[^*]+\*\*[：:]', line):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.first_line_indent = Cm(0.74)
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                add_run_to_paragraph(p, part[2:-2], 10.5, True, None, 'Times New Roman', '宋体')
            else:
                add_run_to_paragraph(p, part, 10.5, False, None, 'Times New Roman', '宋体')

    # Ordered list
    elif re.match(r'^\d+\. ', line):
        text = re.sub(r'^\d+\. ', '', line)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.first_line_indent = Cm(0.74)
        add_run_to_paragraph(p, text, 10.5, False, None, 'Times New Roman', '宋体')

    # Bullet list
    elif line.startswith('- ') or line.startswith('  - '):
        text = re.sub(r'^\s*-\s*', '', line)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.first_line_indent = Cm(0.74)
        add_run_to_paragraph(p, '• ' + text, 10.5, False, None, 'Times New Roman', '宋体')

    # Blockquote
    elif line.startswith('> '):
        text = line[2:]
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(0)
        add_run_to_paragraph(p, text, 10, False, (100, 100, 100), 'Times New Roman', '宋体', italic=True)

    # Empty line
    elif line.strip() == '':
        pass

    # Normal paragraph
    else:
        add_body_p(doc, line)

    i += 1

# ── Save to Desktop ──
output_path = r'C:\Users\Administrator\Desktop\论文-地缘政治压力-v2.docx'
doc.save(output_path)
print(f'Done → {output_path}')
